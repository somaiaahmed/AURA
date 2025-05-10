from docx import Document
import json

def fill_template(template_path, output_path, content):
    document = Document(template_path)

    # Process all paragraphs in the document
    for paragraph in document.paragraphs:
        for placeholder, value in content.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    # Process all tables in the document
    for table in document.tables:
        rows_to_remove = []
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for placeholder, value in content.items():
                    if placeholder in cell.text:
                        # Project Objectives (list)
                        if placeholder == '{Project_Objective_1}' and isinstance(value, list):
                            if value:
                                cell.text = f"1. {value[0]}"
                                for i, obj in enumerate(value[1:], start=2):
                                    new_row = table.add_row()
                                    new_row.cells[0].text = f"{i}. {obj}"

                        # Milestones (list)
                        elif placeholder == '{Milestone_1}' and isinstance(value, list):
                            if value:
                                cell.text = f"1. {value[0]}"
                                for i, milestone in enumerate(value[1:], start=2):
                                    new_row = table.add_row()
                                    new_row.cells[0].text = f"{i}. {milestone}"

                        # Deliverables (2-item list)
                        elif placeholder == '{Project_Deliverable_1}' and isinstance(value, list) and len(value) >= 2:
                            cell.text = cell.text.replace('{Project_Deliverable_1}', value[0])
                            cell.text = cell.text.replace('{Project_Deliverable_2}', value[1])

                        # Stakeholders
                        elif placeholder == '{Name}' and isinstance(value, list) and all(isinstance(item, dict) for item in value):
                            # Mark current row for removal (it contains the placeholder)
                            rows_to_remove.append(row)

                            # Insert stakeholder rows
                            for i, stakeholder in enumerate(value, start=1):
                                new_row = table.add_row()
                                new_row.cells[0].text = f"{i}. {stakeholder['Name']}"
                                new_row.cells[1].text = stakeholder['Role']
                                new_row.cells[2].text = stakeholder['Responsibilities']
                        
                        elif placeholder == "{Phase_1}" and isinstance(value, list):
                            if value:
                                cell.text = f"Phase 1: {value[0]}"
                                for i, phase in enumerate(value[1:], start=2):
                                    new_row = table.add_row()
                                    new_row.cells[0].text = f"Phase{i}: {phase}"

                        # Default/simple replacements
                        else:
                            cell.text = cell.text.replace(placeholder, str(value))

        # Actually remove placeholder rows after processing
        for row in rows_to_remove:
            table._tbl.remove(row._tr)

    document.save(output_path)
    print(f"Document has been successfully filled and saved as {output_path}!")



# Usage
with open('content.json', 'r') as f:
    content_data = json.load(f)

fill_template(
    template_path = r"business-requirement-document-word.docx",
    output_path= r"filled.docx",
    content=content_data
)